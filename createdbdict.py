import os
import random
import datetime
import django
from django.utils import timezone
from django.db import connection, router, transaction, models, DEFAULT_DB_ALIAS
from django.apps import apps
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "healthmanage.settings")
django.setup()
tbtitles=['序号', '字段名', '字段描述', '字段类型', '长度', '主键','自增','外键','为空','默认值']

doc = Document()



# 设置内容

#apps.get_app_config('healthinfo').models
no=0
for m in apps.get_models():
    no+=1
    rowtitles=m._meta.verbose_name+'数据表('+m._meta.db_table+')'
    rowtitles=str(no)+'、'+rowtitles
    doc.add_heading(rowtitles,1)
    print(no,rowtitles)
    table = doc.add_table(rows=1, cols=len(tbtitles))
    # 设置表格样式s
    table.style = 'Light List Accent 1'
    # 设置标题
    title_cells = table.rows[0].cells
    for i in range(len(tbtitles)):
        title_cells[i].text = tbtitles[i]
    title_cells[2].width=10
    j=0
    for f in m._meta.fields:    
        j+=1
        row_cells = table.add_row().cells
        dbtype=''
        isauto='×'
        isnull='×'
        isfk='×'
        fdname=f.name
        maxlength=f.max_length
        if f.primary_key:
            ispk='√'
        else:
            ispk='×'
            
        if f.null:
            isnull='√'
        else:
            isnull='×'
        if f.default==django.db.models.fields.NOT_PROVIDED:
            vdefault=''
        elif hasattr(f.default,'__call__'):
            vdefault='now()'
        elif f.default==None:
            vdefault=''
        else:
            vdefault=f.default

        print(f.get_internal_type())
        if f.get_internal_type() == 'AutoField': 
            dbtype='int'
            isauto='√'
        elif f.get_internal_type() == 'BigAutoField': 
            dbtype='bigint' 
            isauto='√'
        elif f.get_internal_type() == 'ForeignKey': 
            dbtype='ForeignKey'  
            fdname=fdname+'_id'
            dbtype='int'
            isfk='√' 
        elif f.get_internal_type()=='OneToOneField': 
            dbtype='OneToOneField'
            fdname=fdname+'_id'
            isfk='√' 
        elif f.get_internal_type() == 'TextField': 
            dbtype='longtext' 
        elif f.get_internal_type()=='BinaryField': 
            dbtype='longblob' 
        elif f.get_internal_type()=='BooleanField': 
            dbtype='tinyint' 
        elif f.get_internal_type()=='CharField': 
            dbtype='varchar' 
        elif f.get_internal_type()=='CommaSeparatedIntegerField': 
            dbtype='varchar' 
        elif f.get_internal_type()=='DateField': 
            dbtype='date' 
        elif f.get_internal_type()=='DateTimeField': 
            dbtype='datetime' 
        elif f.get_internal_type()=='DecimalField': 
            dbtype='numeric'
            maxlength=str(f.max_digits)+','+str(f.decimal_places)
            print(maxlength)
        elif f.get_internal_type()=='DurationField': 
            dbtype='bigint' 
        elif f.get_internal_type()=='FileField': 
            dbtype='varchar'
        elif f.get_internal_type()=='FilePathField': 
            dbtype='varchar' 
        elif f.get_internal_type()=='FilePathField': 
            dbtype='varchar' 
        elif f.get_internal_type()=='FloatField': 
            dbtype='double' 
        elif f.get_internal_type()=='IntegerField': 
            dbtype='int' 
        elif f.get_internal_type()=='BigIntegerField': 
            dbtype='bigint' 
        elif f.get_internal_type()=='IPAddressField': 
            dbtype='char(15)' 
            maxlength='15'
        elif f.get_internal_type()=='GenericIPAddressField': 
            dbtype='char(39)' 
            maxlength='39'
        elif f.get_internal_type()=='NullBooleanField': 
            dbtype='tinyint' 
        elif f.get_internal_type()=='PositiveIntegerField': 
            dbtype='integer unsigned' 
        elif f.get_internal_type()=='PositiveSmallIntegerField': 
            dbtype='smallint unsigned'  
        elif f.get_internal_type()=='SlugField': 
            dbtype='varchar'  
        elif f.get_internal_type()=='SmallIntegerField': 
            dbtype='smallint' 
        elif f.get_internal_type()=='TimeField': 
            dbtype='time' 
        elif f.get_internal_type()=='UUIDField': 
            dbtype='char(32)'
            maxlength='32'
        else:
            dbtype=f.get_internal_type()
        print(maxlength)
        if maxlength==None:
            maxlength=''
        d=[j,fdname,f.verbose_name,dbtype,maxlength,ispk,isauto,isfk,isnull,vdefault]
        for i in range(len(tbtitles)):
            row_cells[i].text = str(d[i])
doc.save('dict.docx')
print('ok')